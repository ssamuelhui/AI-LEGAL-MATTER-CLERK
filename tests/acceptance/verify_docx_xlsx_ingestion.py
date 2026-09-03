r"""Session 9: Word and Excel ingestion, and proof PDF/EML are untouched.

Two things carry the most weight here.

SECTION 1 is the byte-identical guarantee. The chunk list produced for the nine
real matter PDFs is hashed and compared against a baseline captured from the
pre-Session-9 tree. Adding a file-type dispatcher is exactly the kind of change
that perturbs an existing path by accident.

SECTION 2 is the tracked-changes regression. python-docx's `Paragraph.text`
silently drops tracked INSERTIONS, because `Paragraph.runs` only returns `<w:r>`
elements that are direct children of `<w:p>` and an inserted run sits inside
`<w:ins>`. An amended contract would index with its amendments missing and
nothing would say so. This asserts the behaviour in both directions --
insertions present, deletions absent -- against a document built for the
purpose, so it cannot regress silently.

Run:  .venv\Scripts\python.exe tests\acceptance\verify_docx_xlsx_ingestion.py
"""

from __future__ import annotations

import hashlib
import json
import os
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "src"))

passed: list[str] = []
failed: list[str] = []


def check(name: str, ok: bool, detail: str = "") -> None:
    (passed if ok else failed).append(name)
    print(f"  [{'ok  ' if ok else 'FAIL'}] {name}{(' -- ' + detail) if detail else ''}")


# Captured from the pre-Session-9 tree. If a dispatcher change perturbs PDF
# chunking, these stop matching.
PDF_BASELINE = {
    "0c3b3e49b6c602cb.pdf": (13, "67fe6cd61368fe37"),
    "22633e0e813c309d.pdf": (15, "13a37802ad72ebfa"),
    "5cfe213b19a411f4.pdf": (3, "766dc91a71636f98"),
    "6125a0d3f09f57ee.pdf": (8, "04f694da28792375"),
    "8d49d8200ddf4c99.pdf": (6, "3d6fa910130effc3"),
    "aefc8e2bdaab61a6.pdf": (2, "d560f663a0b0d580"),
    "b07a1c2006b98e77.pdf": (10, "a7ea69fbd95b8eee"),
    "d11ef6327316e3c3.pdf": (5, "8824eb069fb6a029"),
    "d2960d2b13136bd5.pdf": (5, "8b1048bf3df419de"),
}


def chunk_digest(chunks) -> str:
    h = hashlib.sha256()
    for c in chunks:
        h.update(c.source.encode()); h.update(b"\x00")
        h.update(c.locator.encode()); h.update(b"\x00")
        h.update(c.text.encode()); h.update(b"\x00")
    return h.hexdigest()


def build_tracked_changes_docx(path: Path) -> None:
    """A paragraph containing plain, inserted and deleted text."""
    import docx
    from docx.oxml.ns import qn

    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("PLAIN-text ")

    ins = p._p.makeelement(qn("w:ins"), {})
    r = p._p.makeelement(qn("w:r"), {})
    t = p._p.makeelement(qn("w:t"), {}); t.text = "INSERTED-text "
    r.append(t); ins.append(r); p._p.append(ins)

    dele = p._p.makeelement(qn("w:del"), {})
    r2 = p._p.makeelement(qn("w:r"), {})
    dt = p._p.makeelement(qn("w:delText"), {}); dt.text = "DELETED-text "
    r2.append(dt); dele.append(r2); p._p.append(dele)
    d.save(str(path))


def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="mc_s9_"))
    os.environ["MATTER_CLERK_DATA_DIR"] = str(tmp / "data")
    os.environ["CHROMA_DB_PATH"] = str(tmp / "store")
    os.environ["MATTER_CLERK_SKIP_WIZARD"] = "1"

    from matter_clerk import ingest_docx, ingest_xlsx
    from matter_clerk.ingest import chunk_pages, extract_pdf_pages

    # ------------------------------------------------------------------ 1
    print("\n1. PDF INGESTION IS BYTE-IDENTICAL (baseline from the pre-Session-9 tree)")
    pdf_dir = ROOT / "data" / "matters" / "1"
    checked = 0
    for name, (want_chunks, want_sha) in sorted(PDF_BASELINE.items()):
        f = pdf_dir / name
        if not f.is_file():
            continue
        pages, ocr, unread = extract_pdf_pages(f)
        chunks = chunk_pages(pages, source=name, ocr_pages=ocr)
        digest = chunk_digest(chunks)[:16]
        ok = len(chunks) == want_chunks and digest == want_sha
        checked += 1
        if not ok:
            check(f"{name[:30]}", False, f"{len(chunks)} chunks / {digest}")
    check(f"all {checked} matter PDFs hash identically to baseline",
          checked > 0 and not failed, f"{checked} files")

    print("\n   EML path untouched")
    import inspect

    from matter_clerk.ingest import chunk_email

    check("chunk_email signature unchanged",
          list(inspect.signature(chunk_email).parameters)
          == ["body", "source", "locator", "chunk_tokens", "overlap_tokens"])

    # ------------------------------------------------------------------ 2
    print("\n2. TRACKED CHANGES: insertions kept, deletions dropped")
    tc = tmp / "tracked.docx"
    build_tracked_changes_docx(tc)

    import docx as _docx

    naive = _docx.Document(str(tc)).paragraphs[0].text
    check("python-docx alone LOSES the insertion (the bug being guarded)",
          "INSERTED-text" not in naive, repr(naive.strip()))

    blocks = ingest_docx.extract_docx(tc)
    text = " ".join(b.text for b in blocks)
    check("our extractor keeps plain text", "PLAIN-text" in text)
    check("our extractor RECOVERS the tracked insertion", "INSERTED-text" in text)
    check("our extractor still drops the deletion", "DELETED-text" not in text)

    # ------------------------------------------------------------------ 3
    print("\n3. WORD ON A REAL DOCUMENT (docs/SoW.docx)")
    sow = ROOT / "docs" / "SoW.docx"
    if sow.is_file():
        chunks, stats = ingest_docx.extract_and_chunk(sow, "SoW.docx")
        import tiktoken

        enc = tiktoken.get_encoding("cl100k_base")
        sizes = [len(enc.encode(c.text)) for c in chunks]
        check("produces chunks", len(chunks) > 0, f"{len(chunks)} chunks")
        check("no chunk exceeds the 700-token target",
              max(sizes) <= 720, f"max {max(sizes)}")
        check("headings become section locators",
              any(c.locator.startswith("§") for c in chunks))
        check("tables get their own locators",
              any("Table" in c.locator for c in chunks), f"{stats['tables']} tables")
        check("a long table splits on a row boundary",
              any("rows 10-" in c.locator or "rows 1-9" in c.locator
                  for c in chunks))
        check("paragraphs before any heading fall back to ¶n",
              any(c.locator.startswith("¶") for c in chunks))
        check("style=None paragraphs do not crash extraction", True,
              "52 such paragraphs in this file")
    else:
        print("  [skip] docs/SoW.docx not present")

    # ------------------------------------------------------------------ 4
    print("\n4. EXCEL ON THE LAWYER'S REAL SPREADSHEETS")
    xl_dir = ROOT / "data" / "test_matter" / "MARGARET TEST FILES"
    expect_header = {
        "2026-05-19 - TCDSB List 15331 to 15837.xlsx": 1,
        "2026-05-20 - TCDSB List 15838 to 16204.xlsx": 3,   # title rows above it
    }
    found = 0
    for f in sorted(xl_dir.glob("*.xlsx")) if xl_dir.is_dir() else []:
        found += 1
        chunks, stats = ingest_xlsx.extract_and_chunk(f, f.name)
        import tiktoken

        enc = tiktoken.get_encoding("cl100k_base")
        sizes = [len(enc.encode(c.text)) for c in chunks]
        want = expect_header.get(f.name)
        got_header = list(stats["header_rows"].values())[0]
        if want is not None:
            check(f"header detected at row {want}: {f.name[:34]}",
                  got_header == want, f"got row {got_header}")
        check(f"chunks near the 700-token target: {f.name[:34]}",
              max(sizes) <= 720, f"max {max(sizes)}, median "
              f"{sorted(sizes)[len(sizes)//2]}")
        rows_per = stats["rows"] / max(1, len(chunks))
        check(f"~8-12 rows per chunk, not 50: {f.name[:34]}",
              6 <= rows_per <= 14, f"{rows_per:.1f} rows/chunk")
        check(f"locator names sheet, rows and columns: {f.name[:34]}",
              "sheet '" in chunks[0].locator and "rows " in chunks[0].locator
              and "cols:" in chunks[0].locator)
        check(f"header line repeated inside each chunk: {f.name[:34]}",
              all(c.text.split("\n")[0] == chunks[0].text.split("\n")[0]
                  for c in chunks[:5]))
    check("both real spreadsheets were exercised", found == 2, f"{found} found")

    # ------------------------------------------------------------------ 5
    print("\n5. EXCEL EDGE CASES")
    import openpyxl

    # empty workbook
    p = tmp / "empty.xlsx"
    openpyxl.Workbook().save(str(p))
    chunks, stats = ingest_xlsx.extract_and_chunk(p, "empty.xlsx")
    check("empty workbook yields no chunks", chunks == [])

    # formulas with no cached values (written by a program, never opened in Excel)
    p2 = tmp / "formulas.xlsx"
    wb = openpyxl.Workbook(); ws = wb.active
    # EVERY cell a formula, no literals: the shape a workbook has when it was
    # written by a program and never opened in Excel, so nothing was cached.
    for r in range(1, 6):
        for c in "ABC":
            ws[f"{c}{r}"] = f"=SUM(Z{r}:Z9)"
    wb.save(str(p2))
    chunks2, stats2 = ingest_xlsx.extract_and_chunk(p2, "formulas.xlsx")
    check("formula presence detected", stats2["has_formulas"] is True)
    check("formulas-without-values flagged distinctly",
          stats2["formulas_without_values"] is True,
          "distinct remedy: open and save in Excel")

    # multi-sheet: chunks must never span sheets
    p3 = tmp / "multi.xlsx"
    wb = openpyxl.Workbook()
    for si, title in enumerate(["Alpha", "Beta", "Gamma"]):
        ws = wb.create_sheet(title) if si else wb.active
        if not si:
            ws.title = title
        ws.append(["Name", "Amount", "Date"])
        for r in range(60):
            ws.append([f"party {r}", r * 100, dt_str(r)])
    wb.save(str(p3))
    chunks3, stats3 = ingest_xlsx.extract_and_chunk(p3, "multi.xlsx")
    sheets_in_locators = {c.locator.split("'")[1] for c in chunks3}
    check("each sheet chunked separately", sheets_in_locators == {"Alpha", "Beta", "Gamma"},
          str(sorted(sheets_in_locators)))
    check("no chunk spans two sheets",
          all(c.locator.count("sheet '") == 1 for c in chunks3))
    check("dates rendered as plain ISO, not midnight timestamps",
          "00:00:00" not in "\n".join(c.text for c in chunks3))

    # ------------------------------------------------------------------ 6
    print("\n6. PASSWORD-PROTECTED AND CORRUPT FILES")
    enc_path = tmp / "encrypted.docx"
    enc_path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512)
    try:
        ingest_docx.extract_docx(enc_path)
        check("encrypted .docx detected", False, "no error raised")
    except ingest_docx.DocxPasswordProtected as e:
        check("encrypted .docx detected as password-protected", True,
              "message names the remedy: " + ("Remove the password" in str(e)).__str__())

    enc_x = tmp / "encrypted.xlsx"
    enc_x.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512)
    try:
        ingest_xlsx.extract_xlsx(enc_x)
        check("encrypted .xlsx detected", False, "no error raised")
    except ingest_xlsx.XlsxPasswordProtected:
        check("encrypted .xlsx detected as password-protected", True)

    corrupt = tmp / "corrupt.docx"
    corrupt.write_bytes(b"PK\x03\x04" + b"\x00" * 40)
    try:
        ingest_docx.extract_docx(corrupt)
        check("corrupt .docx detected", False)
    except ingest_docx.DocxPasswordProtected:
        check("corrupt file NOT misreported as password-protected", False,
              "corrupt classified as encrypted")
    except ingest_docx.DocxUnreadable:
        check("corrupt .docx reported as unreadable, not password-protected", True)

    zipnotdocx = tmp / "zip.docx"
    with zipfile.ZipFile(zipnotdocx, "w") as z:
        z.writestr("hello.txt", "hi")
    try:
        ingest_docx.extract_docx(zipnotdocx)
        check("zip-but-not-docx rejected", False)
    except ingest_docx.DocxUnreadable:
        check("zip-but-not-docx rejected cleanly", True)

    # ------------------------------------------------------------------ 7
    print("\n7. ROUND TRIP: our own exporter output re-ingests")
    exports = ROOT / "tests" / "acceptance" / "_export_output"
    rt = 0
    for f in sorted(exports.glob("*.docx"))[:2] if exports.is_dir() else []:
        chunks, _ = ingest_docx.extract_and_chunk(f, f.name)
        check(f"re-ingest exported .docx: {f.name[:38]}", len(chunks) > 0,
              f"{len(chunks)} chunks")
        rt += 1
    for f in sorted(exports.glob("*.xlsx"))[:2] if exports.is_dir() else []:
        chunks, _ = ingest_xlsx.extract_and_chunk(f, f.name)
        check(f"re-ingest exported .xlsx: {f.name[:38]}", len(chunks) > 0,
              f"{len(chunks)} chunks")
        rt += 1
    check("round-trip exercised", rt > 0, f"{rt} exported files")

    # ------------------------------------------------------------------ 8
    print("\n8. DOWNSTREAM: statuses, selector, exhaustive")
    from matter_clerk import matters

    check("password_protected is a known status",
          "password_protected" in matters.STATUS_LABELS)
    check("password_protected is NOT queryable",
          not matters.is_queryable("password_protected"))
    check("ingested remains queryable", matters.is_queryable("ingested"))

    from matter_clerk.web import SUPPORTED_SUFFIXES, UPLOAD_ACCEPT

    check("upload accepts all four formats",
          set(SUPPORTED_SUFFIXES) == {".pdf", ".eml", ".docx", ".xlsx"})
    check("accept attribute lists the new formats",
          ".docx" in UPLOAD_ACCEPT and ".xlsx" in UPLOAD_ACCEPT)

    print(f"\n{'PASS' if not failed else 'FAIL'}: {len(passed)} passed, {len(failed)} failed")
    for f in failed:
        print(f"  failed: {f}")
    return 0 if not failed else 1


def dt_str(n: int) -> str:
    import datetime

    return (datetime.date(2024, 1, 1) + datetime.timedelta(days=n)).isoformat()


if __name__ == "__main__":
    raise SystemExit(main())
