"""Low-level Word building blocks (Day 4d).

Every raw python-docx / OOXML detail lives here. The eight task renderers in
`docx_render` compose these and never touch `w:` XML themselves, so a change to
(say) how a table is bordered happens once.
"""

from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .. import pleadings
from .mdblocks import inline_runs
from .payload import PROPERTY_LIMIT, ExportPayload
from .tables import ExportTable

DRAFT_RED = RGBColor(0x7A, 0x00, 0x00)
DRAFT_RED_HEX = "7A0000"
DRAFT_TINT_HEX = "FDF3F3"
HEADER_FILL_HEX = "F2F2F2"
MUTED = RGBColor(0x66, 0x66, 0x66)


# --------------------------------------------------------------------------
# OOXML helpers
# --------------------------------------------------------------------------
def _shade(element_pr, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element_pr.append(shd)


def shade_paragraph(paragraph, fill_hex: str) -> None:
    _shade(paragraph._p.get_or_add_pPr(), fill_hex)


def shade_cell(cell, fill_hex: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), fill_hex)


def set_cell_borders(cell, size: int = 8, color: str = "999999") -> None:
    """size is in eighths of a point (8 == 1pt), per OOXML."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)


def add_page_number_field(paragraph) -> None:
    """Insert a live PAGE field so Word paginates the footer itself."""
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


# --------------------------------------------------------------------------
# Document setup
# --------------------------------------------------------------------------
def apply_base_styles(doc) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(0.9)


def _prop(text: str) -> str:
    """Clamp text to the OOXML core-property limit.

    python-docx raises on any property over 255 characters, which turns an
    export into a failed download. Every property write goes through here so no
    caller-supplied string — a matter name, a task label — can do that again.
    """
    text = str(text)
    if len(text) <= PROPERTY_LIMIT:
        return text
    return text[: PROPERTY_LIMIT - 3].rstrip() + "..."


def set_core_properties(doc, payload: ExportPayload) -> None:
    """Document properties. For a pleading these assert DRAFT status in the
    file's own metadata, so it reads as a draft in a file browser or a document
    management system before anyone opens it.

    Comments carries the SHORT attribution only. The full one names every source
    file and blows the 255-character property limit on any real matter; it is
    written to the page footer instead, where it is visible on every page.
    """
    cp = doc.core_properties
    cp.author = "Matter Clerk (automated draft)"
    cp.comments = _prop(payload.short_attribution())
    if payload.is_pleading:
        cp.title = _prop(f"DRAFT - NOT FOR FILING - {payload.task_label}")
        cp.subject = (
            "DRAFT pleading - not reviewed by counsel - not for filing or service"
        )
        cp.category = "DRAFT"
    else:
        cp.title = _prop(payload.task_label)
        cp.subject = _prop(f"Matter Clerk - {payload.task_label}")


# --------------------------------------------------------------------------
# Blocks
# --------------------------------------------------------------------------
def add_draft_banner(doc) -> None:
    """Full-width dark-red band carrying the code-owned DRAFT_BANNER text.

    A shaded paragraph rather than a table: it survives a lawyer's edits more
    gracefully (no table to accidentally delete as a unit) while remaining
    visually identical to the web UI's `.banner.draft`.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    shade_paragraph(p, DRAFT_RED_HEX)
    run = p.add_run(pleadings.DRAFT_BANNER)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_cover_note_box(doc) -> None:
    """The four SoW 1.4.3 points in a bordered, tinted single-cell table.

    A table (not plain paragraphs) so the note reads as a distinct instrument
    the reader must clear before the body — and so its border survives being
    pasted into another document.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, DRAFT_TINT_HEX)
    set_cell_borders(cell, size=16, color=DRAFT_RED_HEX)  # 2pt red

    intro, *items = pleadings.COVER_NOTE.split("\n1. ", 1)
    body = ("1. " + items[0]) if items else ""

    p = cell.paragraphs[0]
    _write_runs(p, intro.replace("\n", " "))
    for r in p.runs:
        r.font.color.rgb = DRAFT_RED
        r.font.size = Pt(10)

    import re as _re

    for chunk in _re.split(r"\n(?=\d\.\s)", body):
        if not chunk.strip():
            continue
        num, _, rest = chunk.partition(". ")
        item = cell.add_paragraph()
        item.paragraph_format.left_indent = Inches(0.3)
        item.paragraph_format.first_line_indent = Inches(-0.2)
        item.paragraph_format.space_after = Pt(4)
        _write_runs(item, f"{num.strip()}. " + " ".join(rest.split()))
        for r in item.runs:
            r.font.size = Pt(9.5)


def add_authority_disclaimer(doc, payload) -> None:
    """The Phase-2b authority disclaimer, plus this run's verification summary.

    A bordered, tinted single-cell table for the same reason the cover note is
    one: it must read as a distinct instrument and its border must survive being
    pasted into another document. An exported memo can be forwarded to someone
    who never saw the result page, so the file has to carry the caveat itself.
    """
    from .. import verification

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF4E5")
    set_cell_borders(cell, size=16, color="D97706")  # 2pt amber

    p = cell.paragraphs[0]
    run = p.add_run("CITATION VERIFICATION")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x3A, 0x00)

    body = cell.add_paragraph()
    _write_runs(body, verification.AUTHORITY_DISCLAIMER)
    for r in body.runs:
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x6B, 0x3A, 0x00)

    if payload.verification_summary:
        summary = cell.add_paragraph()
        text = payload.verification_summary
        if payload.verification_incomplete:
            text += (
                " — verification did not complete; citations marked "
                "[UNVERIFIED] were not checked."
            )
        run = summary.add_run(f"This draft: {text}.")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x3A, 0x00)
    doc.add_paragraph()


def add_heading(doc, text: str, level: int = 2) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt({1: 15, 2: 13, 3: 11.5}.get(level, 11))


def _write_runs(paragraph, text: str, highlight_markers: bool = False) -> None:
    """Write inline-formatted text. When `highlight_markers` is set, any
    [ELEMENTS REQUIRED ...] / [ADDITIONAL MATERIAL REQUIRED ...] marker is given
    a yellow highlight and bold red type so a reviewer cannot skim past it."""
    segments = (
        pleadings.split_required_markers(text)
        if highlight_markers
        else [(text, False)]
    )
    for segment, is_marker in segments:
        if is_marker:
            run = paragraph.add_run(segment)
            run.bold = True
            run.font.color.rgb = DRAFT_RED
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            continue
        for body, bold, italic, code in inline_runs(segment):
            run = paragraph.add_run(body)
            run.bold = bold
            run.italic = italic
            if code:
                run.font.name = "Consolas"


def add_rich_paragraph(
    doc, text: str, *, highlight_markers: bool = False, size: float | None = None
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    _write_runs(p, text, highlight_markers=highlight_markers)
    if size:
        for r in p.runs:
            r.font.size = Pt(size)
    return p


def add_list(
    doc,
    items: list[str],
    ordered: bool,
    *,
    markers: list[str] | None = None,
    highlight_markers: bool = False,
) -> None:
    """`markers` carries each item's ORIGINAL ordinal. Numbering is reproduced,
    never re-derived — see the note on Block.markers."""
    for n, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(5)
        if ordered:
            ordinal = markers[n] if markers and n < len(markers) else str(n + 1)
            bullet = f"{ordinal}.\t"
        else:
            bullet = "•\t"
        _write_runs(p, bullet, highlight_markers=False)
        _write_runs(p, item, highlight_markers=highlight_markers)


def add_quote(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(7)
    _write_runs(p, text)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_table(doc, table: ExportTable, *, highlight_markers: bool = False) -> None:
    """Render an ExportTable with a bold, shaded header row and full borders."""
    if not table.headers:
        return
    t = doc.add_table(rows=1, cols=len(table.headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, name in enumerate(table.headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        shade_cell(cell, HEADER_FILL_HEX)
        run = cell.paragraphs[0].add_run(name)
        run.bold = True
        run.font.size = Pt(9.5)

    for row in table.rows:
        cells = t.add_row().cells
        for i, value in enumerate(row[: len(table.headers)]):
            para = cells[i].paragraphs[0]
            _write_runs(para, value, highlight_markers=highlight_markers)
            for r in para.runs:
                r.font.size = Pt(9.5)

    if table.note:
        note = doc.add_paragraph()
        note.paragraph_format.space_after = Pt(8)
        run = note.add_run(table.note)
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED


def add_warning_box(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shade_paragraph(p, "FFF8E1")
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x6A, 0x4A, 0x00)


def add_title_block(doc, payload: ExportPayload) -> None:
    add_heading(doc, payload.task_label, level=1)
    bits: list[str] = []
    if payload.matter_name:
        bits.append(f"Matter: {payload.matter_name}")
    if payload.party_role:
        bits.append(f"Party role: {payload.party_role}")
    bits.append("matter-only")
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    run = sub.add_run("  ·  ".join(bits))
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED

    if payload.request_summary:
        for key, value in payload.request_summary.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            k = p.add_run(f"{key}: ")
            k.bold = True
            k.font.size = Pt(9.5)
            v = p.add_run(value)
            v.font.size = Pt(9.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_citations(doc, payload: ExportPayload) -> None:
    """Citations as a numbered block below the content, in the same
    [source locator] form the web UI and the model's inline cites use — so a
    lawyer can match an inline citation to this list by string equality."""
    add_heading(doc, "Citations", level=2)
    if not payload.citations:
        add_rich_paragraph(doc, "_No citations were recorded for this result._", size=9.5)
        return
    if payload.source_files:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"{payload.provenance_label}: " + ", ".join(payload.source_files))
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
    for n, c in enumerate(payload.citations, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(3)
        head = p.add_run(f"{n}. {c.inline()}  ")
        head.bold = True
        head.font.size = Pt(9)
        head.font.name = "Consolas"
        body = p.add_run(f"“{c.text_snippet}”")
        body.font.size = Pt(9)
        body.italic = True


def add_metadata_footer(doc, payload: ExportPayload) -> None:
    """The attribution footer, in the real Word page footer so it appears on
    every page — including any page a lawyer adds while editing."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""

    if payload.is_pleading:
        draft = p.add_run(pleadings.DRAFT_BANNER + "\n")
        draft.bold = True
        draft.font.size = Pt(8)
        draft.font.color.rgb = DRAFT_RED

    run = p.add_run(payload.attribution())
    run.font.size = Pt(7.5)
    run.font.color.rgb = MUTED

    page = footer.add_paragraph()
    page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = page.add_run("Page ")
    label.font.size = Pt(7.5)
    label.font.color.rgb = MUTED
    add_page_number_field(page)
    for r in page.runs:
        r.font.size = Pt(7.5)
        r.font.color.rgb = MUTED
