"""Word generation: a dispatcher over per-task renderers (Day 4d).

The renderers differ because the tasks genuinely differ — Timeline is one table,
Find Entities is a table per category, Draft Pleading is numbered paragraphs
wrapped in safety machinery. They share every drawing primitive, so the
differences stay confined to layout decisions.

THE SAFETY-CRITICAL STRUCTURE: the DRAFT banner and cover note are added by
`build_docx` itself, keyed off `payload.is_pleading` — NOT inside
`render_pleading`. A renderer cannot omit them, and a future edit to the
pleading body renderer cannot drop them. This mirrors the Day-3.5 decision
(ARCHITECTURE 2026-06-10) that the banner and cover note are wrapped around
model output by code rather than requested from the model.
"""

from __future__ import annotations

import io
from typing import Callable

from docx import Document

from . import docx_primitives as prim
from .mdblocks import parse_blocks
from .payload import ExportPayload
from .tables import tables_for


class UnsupportedExport(RuntimeError):
    """Raised when a task has no renderer for the requested format."""


# --------------------------------------------------------------------------
# Shared body walker
# --------------------------------------------------------------------------
def _render_blocks(doc, markdown: str, *, highlight_markers: bool = False) -> None:
    for block in parse_blocks(markdown):
        if block.kind == "heading":
            prim.add_heading(doc, block.text, level=min(block.level + 1, 3))
        elif block.kind == "paragraph":
            prim.add_rich_paragraph(
                doc, block.text, highlight_markers=highlight_markers
            )
        elif block.kind == "list":
            prim.add_list(
                doc,
                block.items,
                block.ordered,
                markers=block.markers,
                highlight_markers=highlight_markers,
            )
        elif block.kind == "quote":
            prim.add_quote(doc, block.text)
        elif block.kind == "code":
            prim.add_rich_paragraph(doc, block.text, size=9)
        elif block.kind == "table" and block.table:
            prim.add_table(doc, block.table, highlight_markers=highlight_markers)


# --------------------------------------------------------------------------
# Per-task renderers
# --------------------------------------------------------------------------
def render_prose(doc, payload: ExportPayload) -> None:
    """Summarize / Find Facts — free-form prose with occasional lists."""
    prim.add_heading(doc, "Answer", level=2)
    _render_blocks(
        doc, payload.answer_markdown,
        highlight_markers=payload.highlight_markers,
    )


def render_sectioned(doc, payload: ExportPayload) -> None:
    """Draft Memo — Issue / Facts / Analysis / Conclusion arrive as markdown
    headings, so the block walker already reproduces the structure; this exists
    to give the body its own titled section and to keep memo-specific layout
    changes away from the plain-prose tasks."""
    prim.add_heading(doc, "Memorandum", level=2)
    # Phase 2b: in authority mode a memo carries [REMOVED — ...] markers
    # recording citations the tool deleted. Those must be as impossible to skim
    # past as a pleading's gap markers, so the memo renderer honours the
    # payload's highlight decision rather than assuming "memos never highlight".
    _render_blocks(
        doc, payload.answer_markdown,
        highlight_markers=payload.highlight_markers,
    )


def render_letter(doc, payload: ExportPayload) -> None:
    """Draft Correspondence — rendered without an 'Answer' heading so the
    letter reads as a letter from the first line."""
    _render_blocks(
        doc, payload.answer_markdown,
        highlight_markers=payload.highlight_markers,
    )


def render_timeline(doc, payload: ExportPayload) -> None:
    tables = tables_for(payload)
    if not tables:
        _render_blocks(doc, payload.answer_markdown)
        return
    for table in tables:
        prim.add_table(doc, table)


def render_entities(doc, payload: ExportPayload) -> None:
    tables = tables_for(payload)
    if not tables:
        _render_blocks(doc, payload.answer_markdown)
        return
    for table in tables:
        if table.title:
            prim.add_heading(doc, table.title, level=3)
        if table.rows:
            prim.add_table(doc, table)
        else:
            prim.add_rich_paragraph(
                doc,
                table.note or "None found in the retrieved passages.",
                size=9.5,
            )


def render_comparison(doc, payload: ExportPayload) -> None:
    """Compare Clauses — one column per document. The section is set landscape
    because a comparison across more than three files is unreadable at portrait
    width, and a crushed table is how a lawyer misreads a cell."""
    tables = tables_for(payload)
    widest = max((len(t.headers) for t in tables), default=0)
    if widest > 3:
        _set_landscape(doc)
    if not tables:
        _render_blocks(doc, payload.answer_markdown)
        return
    for table in tables:
        prim.add_table(doc, table)


def render_pleading(doc, payload: ExportPayload) -> None:
    """Draft Pleading BODY only.

    The DRAFT banner and cover note are deliberately NOT here — `build_docx`
    owns them. This renderer's one special responsibility is highlighting the
    [ELEMENTS REQUIRED ...] / [ADDITIONAL MATERIAL REQUIRED ...] gap markers so
    an unfilled gap is impossible to skim past in a document that may be
    forwarded to a client or opposing counsel.
    """
    prim.add_heading(doc, "Pleading draft", level=2)
    _render_blocks(doc, payload.answer_markdown, highlight_markers=True)


def _set_landscape(doc) -> None:
    from docx.enum.section import WD_ORIENT

    section = doc.sections[0]
    if section.page_width < section.page_height:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )


Renderer = Callable[[object, ExportPayload], None]

DOCX_RENDERERS: dict[str, Renderer] = {
    "summarize": render_prose,
    "find_facts": render_prose,
    "draft_memo": render_sectioned,
    "draft_correspondence": render_letter,
    "timeline": render_timeline,
    "find_entities": render_entities,
    "compare_clauses": render_comparison,
    "draft_pleading": render_pleading,
}


def build_docx(payload: ExportPayload) -> bytes:
    """Assemble the Word document. Order is part of the safety contract."""
    doc = Document()
    prim.apply_base_styles(doc)
    prim.set_core_properties(doc, payload)
    prim.add_metadata_footer(doc, payload)

    renderer = DOCX_RENDERERS.get(payload.task)
    if renderer is None:
        raise UnsupportedExport(f"No Word renderer for task {payload.task!r}.")

    # --- code-owned DRAFT machinery, head -------------------------------
    if payload.is_pleading:
        prim.add_draft_banner(doc)

    prim.add_title_block(doc, payload)

    # Phase 2b: the authority disclaimer precedes the cover note and the body,
    # so a forwarded file carries it before anything a reader might rely on.
    if payload.authority_mode:
        prim.add_authority_disclaimer(doc, payload)

    if payload.is_pleading:
        prim.add_heading(doc, "Cover note", level=2)
        prim.add_cover_note_box(doc)
        doc.add_paragraph()

    for warning in list(payload.pleading_warnings) + list(payload.export_warnings):
        prim.add_warning_box(doc, warning)

    renderer(doc, payload)

    # --- code-owned DRAFT machinery, foot -------------------------------
    if payload.is_pleading:
        prim.add_draft_banner(doc)

    prim.add_citations(doc, payload)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
