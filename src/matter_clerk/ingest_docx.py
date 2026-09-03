r"""Word (.docx) ingestion (Session 9).

Produces the same `Chunk(source, locator, text)` shape as PDF and email, so
everything downstream -- retrieval, the Session 7 file selector, Session 8
exhaustive mode, citations -- works without knowing this format exists.

THREE THINGS THIS DOES THAT python-docx DOES NOT DO FOR YOU
-----------------------------------------------------------

1. TRACKED INSERTIONS ARE RECOVERED. `Paragraph.runs` returns only `<w:r>`
   elements that are DIRECT children of `<w:p>`, and an inserted run lives
   inside `<w:ins>`. So `Paragraph.text` silently drops every tracked
   insertion. Demonstrated during design: a paragraph of
   "plain + inserted + deleted" yields only "plain".

   For a legal tool that is a serious failure -- an amended contract would be
   indexed with its amendments missing and nothing would say so. This module
   walks `w:t` descendants instead, which includes runs inside `w:ins`, and
   skips `w:delText`, which is what deletions use. The result is exactly the
   "all changes accepted" view: insertions in, deletions out.

2. DOCUMENT ORDER IS PRESERVED. `document.paragraphs` and `document.tables`
   are separate collections with no interleaving, so reading them in turn puts
   every table at the end and scrambles the heading a table belongs under. The
   body XML is walked in order instead.

3. MISSING STYLES ARE TOLERATED. 52 of 199 paragraphs in the project's own
   SoW.docx have `style is None`.

COMMENTS are never extracted, and that is free rather than filtered: they live
in `word/comments.xml`, a separate package part that paragraph text never
touches. Privileged review notes stay out of the index by construction.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

import docx
from docx.opc.exceptions import PackageNotFoundError

from .ingest import CHUNK_OVERLAP_TOKENS, CHUNK_TARGET_TOKENS, Chunk, _chunk_tokens

log = logging.getLogger("matter_clerk.ingest_docx")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Deepest heading level tracked in a locator. Beyond this the path gets longer
# than the citation it is meant to label.
MAX_HEADING_DEPTH = 3


class DocxUnreadable(RuntimeError):
    """The file could not be opened as a Word document."""


class DocxPasswordProtected(DocxUnreadable):
    """The file is encrypted."""


@dataclass
class Block:
    """One structural unit of the document, in document order.

    `kind` is "para" or "table". `locator` is the citation label this block's
    text will carry. Blocks are packed into chunks afterwards; keeping the
    split between "what the document contains" and "how it is chunked" is what
    lets the table rules differ from the prose rules without either knowing
    about the other.
    """

    kind: str
    locator: str
    text: str
    heading_path: tuple[str, ...] = ()
    table_index: int = 0
    row_start: int = 0
    row_end: int = 0


def _para_text(p_element) -> str:
    """All text in a paragraph, insertions included, deletions excluded.

    See the module docstring: this is the reason the module exists rather than
    calling `Paragraph.text`.
    """
    parts: list[str] = []
    for node in p_element.iter():
        if node.tag == W + "t":
            parts.append(node.text or "")
        elif node.tag == W + "tab":
            parts.append("\t")
        elif node.tag in (W + "br", W + "cr"):
            parts.append("\n")
    return "".join(parts)


def _style_name(paragraph) -> str:
    style = getattr(paragraph, "style", None)
    return getattr(style, "name", None) or ""


def _heading_level(style_name: str) -> int | None:
    """1/2/3 for Heading 1/2/3, else None. Title counts as level 1."""
    s = (style_name or "").strip().lower()
    if s == "title":
        return 1
    if s.startswith("heading"):
        tail = s[len("heading"):].strip()
        if tail.isdigit():
            return int(tail)
    return None


def _heading_locator(path: tuple[str, ...]) -> str:
    """Render a heading path for a citation.

    Legal documents carry their own numbering in the heading text ("2.3
    Termination"), so nothing is synthesised -- the deepest heading is quoted
    as written, prefixed with a section mark.
    """
    if not path:
        return ""
    return "§" + path[-1]


def open_document(path: Path):
    """Open a .docx, distinguishing encryption from corruption.

    Both raise PackageNotFoundError from python-docx, so the exception cannot
    tell them apart. An encrypted OOXML file is an OLE2 compound document and
    begins with a fixed 8-byte signature; a healthy one is a zip beginning
    "PK". Sniffing the header separates them with no extra dependency.
    """
    try:
        head = path.open("rb").read(8)
    except OSError as e:
        raise DocxUnreadable(f"Could not read {path.name}: {e}") from e

    if head.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
        raise DocxPasswordProtected(
            f"{path.name} is password protected. Remove the password in Word, "
            "save a copy, and upload that copy."
        )
    try:
        return docx.Document(str(path))
    except PackageNotFoundError as e:
        raise DocxUnreadable(
            f"{path.name} could not be opened as a Word document. It may be "
            "damaged, or saved in an older .doc format -- re-save it as .docx."
        ) from e
    except Exception as e:                                        # noqa: BLE001
        raise DocxUnreadable(
            f"{path.name} could not be read as a Word document ({type(e).__name__})."
        ) from e


def extract_docx(path: Path) -> list[Block]:
    """Walk the document body in order, returning blocks with locators."""
    document = open_document(path)

    # Map the underlying XML elements back to python-docx objects so the body
    # can be walked in true document order while still using the library's
    # table and paragraph wrappers.
    paragraphs = {p._p: p for p in document.paragraphs}
    tables = {t._tbl: t for t in document.tables}

    blocks: list[Block] = []
    heading_path: list[str] = []
    para_no = 0
    table_no = 0

    for child in document.element.body.iterchildren():
        if child.tag == W + "p":
            paragraph = paragraphs.get(child)
            if paragraph is None:
                continue
            text = _para_text(child).strip()
            para_no += 1
            if not text:
                continue

            level = _heading_level(_style_name(paragraph))
            if level is not None:
                # Trim the path to this level, then push. A Heading 3 after a
                # Heading 1 with no Heading 2 simply nests one deeper.
                del heading_path[min(level - 1, len(heading_path)):]
                heading_path.append(text)
                heading_path[:] = heading_path[:MAX_HEADING_DEPTH]
                # The heading itself is content: it is often the only place a
                # clause is named, and dropping it loses that name from the index.
                blocks.append(Block(
                    kind="para", text=text, heading_path=tuple(heading_path),
                    locator=_heading_locator(tuple(heading_path))
                             or f"¶{para_no}",
                ))
                continue

            locator = (_heading_locator(tuple(heading_path))
                       or f"¶{para_no}")
            blocks.append(Block(kind="para", text=text,
                                heading_path=tuple(heading_path),
                                locator=locator))

        elif child.tag == W + "tbl":
            table = tables.get(child)
            if table is None:
                continue
            table_no += 1
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = " ".join(
                        _para_text(p._p).strip() for p in cell.paragraphs
                    ).strip()
                    cells.append(cell_text)
                # a row of entirely empty cells carries nothing
                if any(c for c in cells):
                    rows.append(" | ".join(cells))
            if not rows:
                continue
            label = _table_label(table_no, heading_path, rows)
            for i, row_text in enumerate(rows, start=1):
                blocks.append(Block(
                    kind="table", text=row_text,
                    heading_path=tuple(heading_path),
                    table_index=table_no, row_start=i, row_end=i,
                    locator=label,
                ))
    return blocks


def _table_label(index: int, heading_path: list[str], rows: list[str]) -> str:
    """`Table 3 "Fee Schedule"` where a caption is available, else `Table 3`.

    The first row is used as a name only when it looks like a header rather
    than data -- short cells, no long prose.
    """
    first = rows[0] if rows else ""
    cells = [c.strip() for c in first.split("|")]
    looks_like_header = (
        len(cells) >= 2
        and all(len(c) <= 40 for c in cells)
        and any(c for c in cells)
    )
    if looks_like_header:
        name = " / ".join(c for c in cells[:3] if c)[:48]
        if name:
            return f'Table {index} "{name}"'
    return f"Table {index}"


def chunk_docx(
    blocks: list[Block],
    source: str,
    chunk_tokens: int = CHUNK_TARGET_TOKENS,
    overlap_tokens: int = CHUNK_OVERLAP_TOKENS,
) -> list[Chunk]:
    """Pack blocks into chunks, never crossing a heading or table boundary.

    Prose accumulates until the token budget is reached. A single section
    larger than the budget splits into parts that all keep the section's
    locator plus a part marker, so a citation still names the section a lawyer
    can turn to.

    Tables are packed by whole rows only. Half a row of a party schedule is
    worse than no row: the columns stop lining up and the fragment reads as a
    different fact.
    """
    chunks: list[Chunk] = []

    groups: list[tuple[str, list[Block]]] = []
    for block in blocks:
        key = f"{block.kind}:{block.locator}"
        if groups and groups[-1][0] == key:
            groups[-1][1].append(block)
        else:
            groups.append((key, [block]))

    for _key, group in groups:
        locator = group[0].locator
        is_table = group[0].kind == "table"

        if is_table:
            header = group[0].text
            packed: list[str] = []
            running = 0
            part_rows: list[tuple[int, int]] = []
            start_row = group[0].row_start
            for block in group:
                size = _count(block.text)
                if packed and running + size > chunk_tokens:
                    chunks.append(Chunk(
                        source=source,
                        locator=_table_part_locator(locator, start_row,
                                                   block.row_start - 1),
                        text="\n".join(packed),
                    ))
                    # repeat the header so a continuation chunk is readable
                    packed = [header] if header else []
                    running = _count(header) if header else 0
                    start_row = block.row_start
                packed.append(block.text)
                running += size
                part_rows.append((block.row_start, block.row_end))
            if packed:
                chunks.append(Chunk(
                    source=source,
                    locator=_table_part_locator(locator, start_row,
                                               group[-1].row_end),
                    text="\n".join(packed),
                ))
            continue

        text = "\n".join(b.text for b in group)
        pieces = _chunk_tokens(text, chunk_tokens, overlap_tokens)
        if len(pieces) == 1:
            chunks.append(Chunk(source=source, locator=locator, text=pieces[0]))
        else:
            total = len(pieces)
            for i, piece in enumerate(pieces, start=1):
                chunks.append(Chunk(
                    source=source,
                    locator=f"{locator} (part {i} of {total})",
                    text=piece,
                ))
    return chunks


def _table_part_locator(base: str, first_row: int, last_row: int) -> str:
    if last_row <= first_row:
        return f"{base}, row {first_row}"
    return f"{base}, rows {first_row}-{last_row}"


def _count(text: str) -> int:
    from .ingest import _ENC

    return len(_ENC.encode(text or ""))


def extract_and_chunk(path: Path, source: str) -> tuple[list[Chunk], dict]:
    """Full pipeline for one .docx. Returns (chunks, stats)."""
    blocks = extract_docx(path)
    chunks = chunk_docx(blocks, source=source)
    stats = {
        "blocks": len(blocks),
        "paragraph_blocks": sum(1 for b in blocks if b.kind == "para"),
        "table_blocks": sum(1 for b in blocks if b.kind == "table"),
        "tables": len({b.table_index for b in blocks if b.kind == "table"}),
        "headings": len({b.heading_path for b in blocks if b.heading_path}),
        "chars": sum(len(b.text) for b in blocks),
    }
    return chunks, stats
